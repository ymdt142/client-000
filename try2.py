import pandas as pd
from collections import Counter
import docx
df=pd.read_csv("Data.csv")
desgine=docx.Document("Design.docx")
design_dict={}
fail=[]
for content in desgine.paragraphs:
    try:
        design_dict[content.text.split()[0]]=content.text.split()[1]
    except:
        fail.append(content.text)
def findType(user_element): #take user element name as input
    TypeOfElement=design_dict[user_element]
    return TypeOfElement
rules={"SLICE":["LUT1","LUT2","LUT3","LUT4","LUT5","LUT6","FDRE"],"DSP":["DSP48E2"],"RAM":["RAMB36E2"],"IO":["IBUF","OBUF","BUFGCE"]}
def findBelongingTo(TypeOfElement):  #take TypeofElement name as input
    belongsTo=""
    for k,v in rules.items():
        for i in v:
            if i==TypeOfElement:
                belongsTo=k
                return belongsTo
doc=docx.Document("Model.docx")
Model_X=[]
Model_Type=[]
fail=[]
for i in doc.paragraphs[1:len(doc.paragraphs)]:
    try:
        Model_X.append(int(i.text.split()[0]))
        Model_Type.append(i.text.split()[1])
    except:
        fail.append(i.text)
        print(i.text)
def findingXToMove(belongsTo): #get it from findBelongingTo function
    IndexForBelongTo=[]
    for x,belongsto in zip(Model_X,Model_Type):
        if(belongsTo==belongsto):
            IndexForBelongTo.append(x)
    return IndexForBelongTo
#pop once used. Find the difference
def closestTo(user_x,IndexForBelongTo): #get it from findingXToMove
    pre = 999999999
    for index, val in enumerate(IndexForBelongTo):
        curr = (user_x - val) * (user_x - val)
        if (curr == 0):
            return val
        if (curr >= pre):
            return pre_val
        pre = curr
        pre_val = val
    return IndexForBelongTo[-1]
rule_max={"SLICE":30,"DSP":1,"RAM":1,"IO":60}
def findMax(belongsTo): #get it from findBelongingTo function
    maximum = rule_max[belongsTo]
    return maximum
counter={"SLICE":{},"DSP":{},"RAM":{},"IO":{}} # {number1:cnt,number2:cnt}
def checkMax(belongsTo,currentlyUsingIndex,user_y,maximum):  # get belongsTo from findBelongingTo, currentlyUsingIndex from closestTo and maximim from findMax
    currentlyUsingCoordinate=(currentlyUsingIndex,user_y)
    if currentlyUsingCoordinate in counter[belongsTo]:
        for index, cnt in counter[belongsTo].items():
            if index == currentlyUsingCoordinate and cnt == maximum:
                return True
        counter[belongsTo][currentlyUsingCoordinate]=counter[belongsTo][currentlyUsingCoordinate]+1
        return False
    else:
        counter[belongsTo][currentlyUsingCoordinate] = 1
        return False


# taking input ony by one from Data
def write(index,msg):
    files=["log1.txt","log2.txt","log3.txt"]
    if index<30000:
        with open(files[0], "a+") as obj1:
            obj1.write(msg)
    if index>=30000 and index <=60000:
        with open(files[1], "a+") as obj2:
            obj2.write(msg)
    if index >60000:
        with open(files[2], "a+") as obj3:
            obj3.write(msg)

usedY={}
def set_Y(y,usedY,noX=False):
    if y not in usedY:
        usedY[y]=0
        return y
    if noX:
        usedY[y]=1
        for y_axis in range(y,481):
            if y_axis in usedY:
                if usedY[y_axis]==0:
                    return y_axis
            else:
                usedY[y_axis]=0
                return y_axis
        for y_axis in range(y,1,-1):
            if y_axis in usedY:
                if usedY[y_axis]==0:
                    return y_axis
            else:
                usedY[y_axis]=0
                return y_axis
    if usedY[y]==1:
        for y_axis in range(y,481):
            if y_axis in usedY:
                if usedY[y_axis]==0:
                    return y_axis
            else:
                usedY[y_axis]=0
                return y_axis
        for y_axis in range(y,1,-1):
            if y_axis in usedY:
                if usedY[y_axis]==0:
                    return y_axis
            else:
                usedY[y_axis]=0
                return y_axis
    if usedY[y]==0:
        return y

def fixed(filename):
    d = docx.Document(filename)
    dont_change = []
    for c in d.paragraphs:
        try:
            element = c.text.split()[0] + " " + c.text.split()[1] + " " + c.text.split()[2]
            dont_change.append(element)
        except:
            pass
    return dont_change

dont_change=fixed("design1.docx")
new_ele = {}
# taking input ony by one from Data

dont_use = []
for index, row in df.iterrows():
    user_element = row["Element"]
    user_x = int(row["X"])
    true_user_Y = int(row["y"])
    name=user_element+" "+str(user_x)+" "+str(true_user_Y)
    if name not in dont_change:
        user_y = set_Y(int(row["y"]), usedY)
        msg = "Index" + " " + str(index) + " " + user_element + " " + str(user_x) + " " + str(true_user_Y) + "\n"
        write(index, msg)
        TypeOfElement = findType(user_element)  # find type
        belongsTo = findBelongingTo(TypeOfElement)  # finding belongs to
        IndexForBelongTo = findingXToMove(belongsTo)  # finding x to move element from model
        for du in dont_use:
            try:
                IndexForBelongTo.remove(du[0])
            except:
                pass  # assigining Index for new Y
        if (len(IndexForBelongTo) == 0):
            user_y = set_Y(int(row["y"]), usedY, True)
            IndexForBelongTo = findingXToMove(belongsTo)
        closest = closestTo(user_x, IndexForBelongTo)  # finding cloest to Index
        maximum = findMax(belongsTo)  # finding max for certain type element
        old = user_element + " " + str(user_x) + " " + str(true_user_Y)
        if (not checkMax(belongsTo, closest, user_y, maximum)):
            new_ele[old] = belongsTo + " " + user_element + " " + str(closest) + " " + str(user_y)
        else:
            dont_use.append((closest, user_y))

            for du in dont_use:
                try:
                    IndexForBelongTo.remove(du[0])
                except:
                    pass
            if (len(IndexForBelongTo) == 0):
                user_y = set_Y(int(row["y"]), usedY, True)
                IndexForBelongTo = findingXToMove(belongsTo)
            closest = closestTo(user_x, IndexForBelongTo)  # finding cloest to Index
            checkMax(belongsTo, closest, user_y, maximum)  # putting the number to counter
            new_ele[old] = belongsTo + " " + user_element + " " + str(closest) + " " + str(user_y)
        op = ["output1.csv", "output2.csv", "output3.csv"]
        if index == 29999:
            op1 = pd.DataFrame(pd.Series(new_ele))
            op1.to_csv(op[0])
            new_ele = {}
            try:
                geeky_file = open("counter1.txt", 'wt')
                geeky_file.write(str(counter))
                geeky_file.close()

            except:
                print("Unable to write to file")
        if index == 59999:
            op2 = pd.DataFrame(pd.Series(new_ele))
            op2.to_csv(op[1])
            new_ele = {}
            try:
                geeky_file = open("counter2.txt", 'wt')
                geeky_file.write(str(counter))
                geeky_file.close()

            except:
                print("Unable to write to file")
        if index == df.shape[0] - 1:
            op3 = pd.DataFrame(pd.Series(new_ele))
            op3.to_csv(op[2])
            new_ele = {}
            try:
                geeky_file = open("counter3.txt", 'wt')
                geeky_file.write(str(counter))
                geeky_file.close()

            except:
                print("Unable to write to file")
    else:
        print("Dont Change name  ",name)






